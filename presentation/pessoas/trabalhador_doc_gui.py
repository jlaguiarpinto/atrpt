# atrpt/presentation/pessoas/trabalhador_doc_gui.py

import re
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from pathlib import Path
from presentation.shared.base_gui import BaseGui as BG
import logging

_PH_RE = re.compile(r'\{\{[^}]+\}\}')

logger = logging.getLogger(__name__)

_SIT_LABEL = {
    'A': 'Ativo', 'I': 'Inativo', 'B': 'Baixa',
    'L': 'Licença', 'S': 'Suspensão', 'P': 'Pré-reforma', 'C': 'Candidato',
}

# Nomes fixos dos ficheiros de minuta na pasta minutas_dir
_MINUTA_CONTRATO = "Minuta_Contrato_sem termo.docx"
_MINUTA_REGISTO  = "Registo_biográfico_profissional_minuta.docx"


def _fmt_date(d):
    return d.strftime("%d/%m/%Y") if d else ""


def _fmt_float(v):
    return f"{v:.2f}" if v is not None else ""


def _build_mapping(e) -> dict:
    return {
        # identificação
        "{{nome}}":                  e.nome or "",
        "{{nome_conhecido}}":        e.nome_conhecido or "",
        "{{estado_civil}}":          e.estado_civil or "",
        "{{estado:civil}}":          e.estado_civil or "",   # alias com : no registo
        "{{nacionalidade}}":         e.nacionalidade or "",
        "{{data_nascimento}}":       _fmt_date(e.data_nascimento),
        "{{id_tipo}}":               e.doc_ident_tipo or "",
        "{{id_num}}":                e.cc or "",
        "{{id_val}}":                _fmt_date(e.data_validade_cc),
        "{{id_data}}":               _fmt_date(e.data_validade_cc),
        "{{NIF}}":                   e.nif or "",
        "{{NISS}}":                  e.niss or "",
        "{{n_empregado}}":           str(e.numero),
        # contacto
        "{{morada}}":                e.morada or "",
        "{{localidade}}":            e.localidade or "",
        "{{cod_postal}}":            e.cp or "",
        "{{telemovel}}":             e.telemovel or "",
        "{{email}}":                 e.email or "",
        "{{EMAIL}}":                 e.email or "",          # alias maiúsculas
        # emergência
        "{{nome_emerg}}":            e.emerg_nome or "",
        "{{parente_emeg}}":          e.emerg_relacao or "",
        "{{telemovelemerg}}":        e.emerg_telemovel or "",
        # bancário
        "{{banco}}":                 e.banco or "",
        "{{IBAN}}":                  e.nib or "",
        # contrato
        "{{categoria_profissional}}": e.categoria_atual or "",
        "{{data_início}}":           _fmt_date(e.data_admissao),
        "{{data_inicio}}":           _fmt_date(e.data_admissao),
        "{{retribuicao}}":           _fmt_float(e.vencimento),
    }


def _placeholders_em_doc(doc) -> set:
    found = set()
    def _scan(para):
        found.update(_PH_RE.findall("".join(r.text for r in para.runs)))
    for p in doc.paragraphs:
        _scan(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _scan(p)
    return found


def _replace_in_doc(doc, mapping: dict) -> None:
    def _replace_para(para):
        for run in para.runs:
            for ph, val in mapping.items():
                if ph in run.text:
                    run.text = run.text.replace(ph, val)
        full = "".join(r.text for r in para.runs)
        if any(ph in full for ph in mapping):
            for ph, val in mapping.items():
                full = full.replace(ph, val)
            if para.runs:
                para.runs[0].text = full
                for r in para.runs[1:]:
                    r.text = ""

    for para in doc.paragraphs:
        _replace_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_para(para)


class TrabalhadorDocGUI(BG):
    """Gera os documentos de admissão a partir das minutas configuradas."""

    def __init__(self, parent, controller):
        super().__init__(parent, controller)
        self._empregado = None

    def _post_init(self):
        self.build_menu_buttons([])
        self._build()

    def _build(self):
        outer = tk.Frame(self.frame, bg=self.BG)
        outer.pack(fill="both", expand=True, padx=10, pady=8)
        outer.columnconfigure(1, weight=1)

        # ── pesquisa ──────────────────────────────────────────────────
        tk.Label(outer, text="Pesquisar trabalhador:", bg=self.BG, fg=self.FG,
                 font=self.FONT_BUTTON).grid(row=0, column=0, sticky="w", pady=(0, 4))

        search_frm = tk.Frame(outer, bg=self.BG)
        search_frm.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 6))

        self._v_pesquisa = tk.StringVar()
        ttk.Entry(search_frm, textvariable=self._v_pesquisa, width=40).pack(side="left", padx=(0, 6))
        tk.Button(search_frm, text="Pesquisar", command=self._pesquisar,
                  font=self.FONT_BUTTON, bg=self.BTN_BG).pack(side="left")
        search_frm.winfo_children()[0].bind("<Return>", lambda e: self._pesquisar())

        # ── lista de resultados ────────────────────────────────────────
        list_frm = tk.Frame(outer, bg=self.BG)
        list_frm.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(0, 8))
        list_frm.columnconfigure(0, weight=1)

        cols = ("Num", "Nome", "Situação", "Local")
        self._tree = ttk.Treeview(list_frm, columns=cols, show="headings", height=6)
        self._tree.heading("Num",      text="Nº")
        self._tree.heading("Nome",     text="Nome")
        self._tree.heading("Situação", text="Situação")
        self._tree.heading("Local",    text="Local")
        self._tree.column("Num",      width=50,  anchor="w")
        self._tree.column("Nome",     width=240, anchor="w")
        self._tree.column("Situação", width=90,  anchor="w")
        self._tree.column("Local",    width=120, anchor="w")

        vsb = ttk.Scrollbar(list_frm, orient="vertical", command=self._tree.yview)
        self._tree.configure(yscrollcommand=vsb.set)
        self._tree.grid(row=0, column=0, sticky="ew")
        vsb.grid(row=0, column=1, sticky="ns")
        self._tree.bind("<<TreeviewSelect>>", self._on_select)
        self._iid_map = {}

        # ── ficha resumida ─────────────────────────────────────────────
        self._lbl_ficha = tk.Label(outer, text="Nenhum trabalhador seleccionado.",
                                   bg=self.BG, fg="gray", justify="left",
                                   font=("Segoe UI", 9))
        self._lbl_ficha.grid(row=3, column=0, columnspan=2, sticky="w", pady=(0, 10))

        # ── minutas ───────────────────────────────────────────────────
        ttk.Separator(outer, orient="horizontal").grid(
            row=4, column=0, columnspan=2, sticky="ew", pady=4)

        minutas_dir = self._minutas_dir()
        info_txt = (
            f"Minutas:  {minutas_dir}\n"
            f"  • {_MINUTA_CONTRATO}\n"
            f"  • {_MINUTA_REGISTO}"
        )
        self._lbl_minutas = tk.Label(outer, text=info_txt, bg=self.BG, fg=self.FG,
                                     font=("Segoe UI", 9), justify="left")
        self._lbl_minutas.grid(row=5, column=0, columnspan=2, sticky="w", pady=(2, 6))

        # ── pasta de destino ──────────────────────────────────────────
        tk.Label(outer, text="Pasta de destino:", bg=self.BG, fg=self.FG).grid(
            row=6, column=0, sticky="w", pady=3)
        dest_frm = tk.Frame(outer, bg=self.BG)
        dest_frm.grid(row=6, column=1, sticky="ew", pady=3)
        self._v_destino = tk.StringVar(value=str(self._admissoes_dir()))
        ttk.Entry(dest_frm, textvariable=self._v_destino, width=54).pack(side="left", padx=(0, 4))
        tk.Button(dest_frm, text="...", command=self._browse_destino,
                  font=("Segoe UI", 9), bg=self.BTN_BG, width=3).pack(side="left")

        # ── pré-visualização dos nomes ────────────────────────────────
        self._lbl_ficheiros = tk.Label(outer, text="", bg=self.BG, fg="gray",
                                       font=("Segoe UI", 9), justify="left")
        self._lbl_ficheiros.grid(row=7, column=0, columnspan=2, sticky="w", pady=(2, 4))

        # ── botão gerar ───────────────────────────────────────────────
        tk.Button(outer, text="Gerar Documentos", command=self._gerar,
                  font=self.FONT_BUTTON, bg=self.BTN_BG).grid(
            row=8, column=0, columnspan=2, pady=14)

        self._carregar_todos()

    # ------------------------------------------------------------------

    def _minutas_dir(self) -> Path:
        p = self.controller.cfg.paths.get("minutas_dir")
        if p and Path(p).exists():
            return Path(p)
        # fallback: subpasta minutas dentro de admissões
        return self._admissoes_dir() / "minutas"

    def _admissoes_dir(self) -> Path:
        p = self.controller.cfg.paths.get("admissoes_dir")
        if p and Path(p).exists():
            return Path(p)
        return Path.home() / "Documents"

    def _nome_ficheiro(self, e) -> str:
        return "".join(c if c.isalnum() or c in " _-" else "_" for c in e.nome).strip()

    def _actualizar_preview(self):
        if self._empregado:
            pasta = self._v_destino.get().strip() or str(self._admissoes_dir())
            nf = self._nome_ficheiro(self._empregado)
            self._lbl_ficheiros.config(
                text=f"Ficheiros a criar em  {pasta}:\n"
                     f"  • {nf}_contrato.docx\n"
                     f"  • {nf}_registo_biografico.docx",
                fg=self.FG,
            )
        else:
            self._lbl_ficheiros.config(text="", fg="gray")

    def _carregar_todos(self):
        dados = self.controller.get_trabalhadores(situacao=None)
        self._popular(dados)

    def _pesquisar(self):
        txt = self._v_pesquisa.get().strip()
        if not txt:
            self._carregar_todos()
            return
        dados = self.controller.pesquisar(txt, apenas_ativos=False)
        self._popular(dados)

    def _popular(self, dados):
        self._tree.delete(*self._tree.get_children())
        self._iid_map.clear()
        for e in dados:
            iid = self._tree.insert("", "end", values=(
                e.numero, e.nome,
                _SIT_LABEL.get(e.ativo, e.ativo),
                e.local or "",
            ))
            self._iid_map[iid] = e

    def _on_select(self, event=None):
        sel = self._tree.selection()
        if not sel:
            self._empregado = None
            self._lbl_ficha.config(text="Nenhum trabalhador seleccionado.", fg="gray")
            self._lbl_ficheiros.config(text="", fg="gray")
            return
        e = self._iid_map.get(sel[0])
        self._empregado = e
        if e:
            linhas = [
                f"N.º {e.numero}  —  {e.nome}",
                f"NIF: {e.nif or '—'}   NISS: {e.niss or '—'}   IBAN: {e.nib or '—'}",
                f"Categoria: {e.categoria_atual or '—'}   Admissão: {_fmt_date(e.data_admissao)}",
            ]
            self._lbl_ficha.config(text="\n".join(linhas), fg=self.FG)
        self._actualizar_preview()

    def _browse_destino(self):
        path = filedialog.askdirectory(
            title="Seleccionar pasta de destino",
            initialdir=self._v_destino.get() or str(self._admissoes_dir()),
            parent=self.root,
        )
        if path:
            self._v_destino.set(path)
            self._actualizar_preview()

    def _gerar(self):
        if not self._empregado:
            messagebox.showwarning("Atenção", "Seleccione um trabalhador.", parent=self.root)
            return

        pasta_dest = self._v_destino.get().strip()
        if not pasta_dest:
            messagebox.showwarning("Atenção", "Indique a pasta de destino.", parent=self.root)
            return
        pasta_dest = Path(pasta_dest)

        minutas = self._minutas_dir()
        contrato_tmpl = minutas / _MINUTA_CONTRATO
        registo_tmpl  = minutas / _MINUTA_REGISTO

        for tmpl, label in [(contrato_tmpl, "contrato"), (registo_tmpl, "registo biográfico")]:
            if not tmpl.exists():
                messagebox.showerror(
                    "Minuta não encontrada",
                    f"O ficheiro de minuta do {label} não foi encontrado:\n{tmpl}",
                    parent=self.root,
                )
                return

        try:
            from docx import Document
        except ImportError:
            messagebox.showerror(
                "Módulo em falta",
                "A biblioteca python-docx não está instalada.\nExecute: pip install python-docx",
                parent=self.root,
            )
            return

        pasta_dest.mkdir(parents=True, exist_ok=True)
        nf = self._nome_ficheiro(self._empregado)
        destino_contrato = pasta_dest / f"{nf}_contrato.docx"
        destino_registo  = pasta_dest / f"{nf}_registo_biografico.docx"

        mapping = _build_mapping(self._empregado)
        avisos_todos = []

        try:
            for tmpl, destino, label in [
                (contrato_tmpl, destino_contrato, "Contrato"),
                (registo_tmpl,  destino_registo,  "Registo biográfico"),
            ]:
                doc = Document(tmpl)
                m = dict(mapping)

                phs_doc     = _placeholders_em_doc(doc)
                em_falta    = sorted(ph for ph in phs_doc if ph in m and not m[ph])
                nao_mapeados = sorted(ph for ph in phs_doc if ph not in m)
                for ph in nao_mapeados:
                    m[ph] = ""

                if em_falta or nao_mapeados:
                    avisos_todos.append(
                        f"{label}:\n" +
                        "\n".join(f"  {ph}" for ph in em_falta + nao_mapeados)
                    )

                _replace_in_doc(doc, m)
                doc.save(destino)
                logger.info("Documento gerado: %s", destino)

        except Exception as ex:
            logger.exception("Erro ao gerar documentos")
            messagebox.showerror("Erro", str(ex), parent=self.root)
            return

        if avisos_todos:
            messagebox.showwarning(
                "Campos sem valor",
                "Os seguintes campos ficam em branco no documento:\n\n" +
                "\n\n".join(avisos_todos),
                parent=self.root,
            )

        messagebox.showinfo(
            "Documentos gerados",
            f"Ficheiros guardados em:\n{pasta_dest}\n\n"
            f"• {destino_contrato.name}\n"
            f"• {destino_registo.name}",
            parent=self.root,
        )
