# presentation/secretaria/contrato_residente_view.py
#
# Gera o contrato de admissão de um residente a partir de uma minuta Word.
# Padrão idêntico ao trabalhador_doc_gui.py (pessoas).
# Acesso: botão em NovoResidenteView  OU  directamente pelo menu Residentes.

import re
import logging
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from pathlib import Path

from presentation.shared.base_gui import BaseGui as BG

logger = logging.getLogger(__name__)

_PH_RE = re.compile(r'\{\{[^}]+\}\}')
_MINUTA_CONTRATO = "Minuta_Contrato_Residente.docx"


# ------------------------------------------------------------------
# Helpers (copiados do padrão trabalhador_doc_gui)
# ------------------------------------------------------------------

def _fmt_date(v) -> str:
    if not v:
        return ""
    s = str(v).strip()
    # aceitar "YYYY-MM-DD" e "DD-MM-YYYY" / "DD/MM/YYYY"
    for sep in ("-", "/", "."):
        parts = s.split(sep)
        if len(parts) == 3:
            if len(parts[0]) == 4:          # YYYY-MM-DD → DD/MM/YYYY
                return f"{parts[2]}/{parts[1]}/{parts[0]}"
            if len(parts[2]) == 4:          # DD-MM-YYYY → DD/MM/YYYY
                return f"{parts[0]}/{parts[1]}/{parts[2]}"
    return s


def _build_mapping(d: dict) -> dict:
    """Cria o dicionário {{placeholder}}: valor a partir dos dados do residente."""
    def v(campo):
        return str(d.get(campo) or "").strip()

    return {
        "{{numero_residente}}": v("numero_residente"),
        "{{numero_socio}}":     v("numero_socio"),
        "{{nome}}":             v("nome"),
        "{{Nome}}":             v("nome"),
        "{{NOME}}":             v("nome"),
        "{{nif}}":              v("nif"),
        "{{NIF}}":              v("nif"),
        "{{genero}}":           v("genero"),
        "{{relacao}}":          v("relacao"),
        "{{petit_nom}}":        v("petit_nom"),
        "{{resp_gen}}":         v("resp_gen"),
        "{{responsavel}}":      v("responsavel"),
        "{{Responsavel}}":      v("responsavel"),
        "{{email}}":            v("email"),
        "{{EMAIL}}":            v("email"),
        "{{resp_tlm}}":         v("resp_tlm"),
        "{{telemovel}}":        v("resp_tlm"),
        "{{iban}}":             v("iban"),
        "{{IBAN}}":             v("iban"),
        "{{data_iban}}":        _fmt_date(v("data_iban")),
        "{{designacao_bancaria}}": v("designacao_bancaria"),
        "{{mensalidade}}":      v("mensalidade"),
        "{{especial}}":         v("especial"),
        "{{copag}}":            v("copag"),
        "{{data_nascimento}}":  _fmt_date(v("data_nascimento")),
        "{{data_admissao}}":    _fmt_date(v("data_admissao")),
        "{{data_inicio}}":      _fmt_date(v("data_admissao")),
        "{{data_início}}":      _fmt_date(v("data_admissao")),
        "{{data_fim}}":         _fmt_date(v("data_fim")),
        # campos de candidatos
        "{{morada}}":           v("morada"),
        "{{cod_postal}}":       v("cod_postal"),
        "{{contato}}":          v("contato"),
        "{{id_tipo}}":          v("id_tipo"),
        "{{id_num}}":           v("id_num"),
        "{{id_val}}":           _fmt_date(v("id_val")),
        "{{resp_id_tipo}}":     v("resp_id_tipo"),
        "{{resp_id_num}}":      v("resp_id_num"),
        "{{resp_id_val}}":      _fmt_date(v("resp_id_val")),
        "{{caucao}}":           v("caucao"),
        "{{mensalidade_contrato}}": v("mensalidade"),
        "{{estado}}":           v("estado"),
    }


def _placeholders_em_doc(doc) -> set:
    found = set()
    def _scan(para):
        found.update(_PH_RE.findall("".join(r.text for r in para.runs)))
    for p in doc.paragraphs:
        _scan(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _scan(p)
    return found


def _replace_in_doc(doc, mapping: dict) -> None:
    def _replace_para(para):
        for run in para.runs:
            for ph, val in mapping.items():
                if ph in run.text:
                    run.text = run.text.replace(ph, val)
        full = "".join(r.text for r in para.runs)
        if any(ph in full for ph in mapping):
            for ph, val in mapping.items():
                full = full.replace(ph, val)
            if para.runs:
                para.runs[0].text = full
                for r in para.runs[1:]:
                    r.text = ""
    for para in doc.paragraphs:
        _replace_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_para(para)


# ------------------------------------------------------------------
# Vista
# ------------------------------------------------------------------

class ContratoResidenteView(BG):
    """Gera o contrato de admissão de um residente a partir de minuta Word."""

    def __init__(self, parent, controller,
                 pre_numero: int | None = None,
                 candidato: dict | None = None):
        super().__init__(parent, controller)
        self._dados: dict | None = None
        self._pre_numero = pre_numero
        self._candidato  = candidato    # dict de candidatos_repo (pré-seleccionado)

    def _post_init(self):
        root = getattr(self, "_root_gui", self)
        self.build_menu_buttons([("← Residentes", root.go_back)])
        self._build()

    # ------------------------------------------------------------------
    # Layout
    # ------------------------------------------------------------------

    def _build(self):
        outer = tk.Frame(self.frame_work, bg=self.BG)
        outer.pack(fill="both", expand=True, padx=10, pady=8)
        outer.columnconfigure(1, weight=1)

        # ── pesquisa / filtro ─────────────────────────────────────────
        tk.Label(outer, text="Pesquisar residente:", bg=self.BG, fg=self.FG,
                 font=self.FONT_BUTTON).grid(row=0, column=0, sticky="w", pady=(0, 4))

        sf = tk.Frame(outer, bg=self.BG)
        sf.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 6))

        self._v_pesq = tk.StringVar()
        ttk.Entry(sf, textvariable=self._v_pesq, width=36).pack(side="left", padx=(0, 6))
        sf.winfo_children()[0].bind("<Return>", lambda _: self._pesquisar())
        tk.Button(sf, text="Pesquisar", command=self._pesquisar,
                  font=("Segoe UI", 9), bg=self.BTN_BG).pack(side="left", padx=(0, 10))

        self._v_so_pendentes = tk.BooleanVar(value=True)
        ttk.Checkbutton(sf, text="Só activos (pendente/espera)", variable=self._v_so_pendentes,
                        command=self._pesquisar).pack(side="left")

        # ── lista ─────────────────────────────────────────────────────
        lf = tk.Frame(outer, bg=self.BG)
        lf.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(0, 8))
        lf.columnconfigure(0, weight=1)

        cols = ("nr", "nome", "especial", "admissao")
        self._tree = ttk.Treeview(lf, columns=cols, show="headings", height=7)
        self._tree.heading("nr",      text="Nº")
        self._tree.heading("nome",    text="Nome")
        self._tree.heading("especial",text="Especial")
        self._tree.heading("admissao",text="Dt. Admissão")
        self._tree.column("nr",       width=50,  anchor="center")
        self._tree.column("nome",     width=240, anchor="w")
        self._tree.column("especial", width=90,  anchor="center")
        self._tree.column("admissao", width=100, anchor="center")

        vsb = ttk.Scrollbar(lf, orient="vertical", command=self._tree.yview)
        self._tree.configure(yscrollcommand=vsb.set)
        self._tree.grid(row=0, column=0, sticky="ew")
        vsb.grid(row=0, column=1, sticky="ns")
        self._tree.bind("<<TreeviewSelect>>", self._on_select)
        self._iid_map: dict[str, dict] = {}

        # ── ficha resumida ─────────────────────────────────────────────
        self._lbl_ficha = tk.Label(
            outer, text="Nenhum residente seleccionado.",
            bg=self.BG, fg="gray", justify="left", font=("Segoe UI", 9),
        )
        self._lbl_ficha.grid(row=3, column=0, columnspan=2, sticky="w", pady=(0, 8))

        ttk.Separator(outer, orient="horizontal").grid(
            row=4, column=0, columnspan=2, sticky="ew", pady=4)

        # ── minuta (seleccionável) ────────────────────────────────────
        tk.Label(outer, text="Minuta (template):", bg=self.BG, fg=self.FG).grid(
            row=5, column=0, sticky="w", pady=3)
        mf = tk.Frame(outer, bg=self.BG)
        mf.grid(row=5, column=1, sticky="ew", pady=3)
        self._v_minuta = tk.StringVar()
        self._cb_minuta = ttk.Combobox(mf, textvariable=self._v_minuta, width=46, state="readonly")
        self._cb_minuta.pack(side="left", padx=(0, 4))
        tk.Button(mf, text="...", command=self._browse_minuta,
                  font=("Segoe UI", 9), bg=self.BTN_BG, width=3).pack(side="left")
        self._popular_minutas()

        # ── destino ───────────────────────────────────────────────────
        tk.Label(outer, text="Pasta de destino:", bg=self.BG, fg=self.FG).grid(
            row=6, column=0, sticky="w", pady=3)
        df = tk.Frame(outer, bg=self.BG)
        df.grid(row=6, column=1, sticky="ew", pady=3)
        self._v_destino = tk.StringVar(value=str(self._admissoes_dir()))
        ttk.Entry(df, textvariable=self._v_destino, width=52).pack(side="left", padx=(0, 4))
        tk.Button(df, text="...", command=self._browse_destino,
                  font=("Segoe UI", 9), bg=self.BTN_BG, width=3).pack(side="left")

        # ── preview nome ficheiro ─────────────────────────────────────
        self._lbl_preview = tk.Label(outer, text="", bg=self.BG, fg="gray",
                                     font=("Segoe UI", 9), justify="left")
        self._lbl_preview.grid(row=7, column=0, columnspan=2, sticky="w", pady=(2, 4))

        # ── botão gerar ───────────────────────────────────────────────
        tk.Button(outer, text="Gerar Contrato", command=self._gerar,
                  font=self.FONT_BUTTON, bg="#c8e6c9", width=20).grid(
            row=8, column=0, columnspan=2, pady=14)

        self._carregar()

    # ------------------------------------------------------------------
    # Paths
    # ------------------------------------------------------------------

    def _minutas_dir(self) -> Path:
        p = self.controller.cfg.paths.get("minutas_residentes_dir")
        if p and Path(p).exists():
            return Path(p)
        return self._admissoes_dir() / "minutas"

    def _popular_minutas(self):
        pasta = self._minutas_dir()
        docx = sorted(pasta.glob("*.docx")) if pasta.exists() else []
        nomes = [f.name for f in docx]
        self._cb_minuta["values"] = nomes
        # pré-seleccionar o template padrão se existir, ou o primeiro disponível
        if _MINUTA_CONTRATO in nomes:
            self._v_minuta.set(_MINUTA_CONTRATO)
        elif nomes:
            self._v_minuta.set(nomes[0])

    def _browse_minuta(self):
        path = filedialog.askopenfilename(
            title="Seleccionar minuta",
            initialdir=str(self._minutas_dir()),
            filetypes=[("Word", "*.docx"), ("Todos", "*.*")],
            parent=self.root,
        )
        if path:
            p = Path(path)
            # acrescentar ao combobox se não estiver na lista
            nomes = list(self._cb_minuta["values"])
            if p.name not in nomes:
                nomes.append(p.name)
                self._cb_minuta["values"] = nomes
                # guardar caminho completo para ficheiros fora da pasta padrão
                self._minutas_extras = getattr(self, "_minutas_extras", {})
                self._minutas_extras[p.name] = p
            self._v_minuta.set(p.name)

    def _admissoes_dir(self) -> Path:
        p = self.controller.cfg.paths.get("admissoes_residentes_dir")
        if p and Path(p).exists():
            return Path(p)
        p2 = self.controller.cfg.paths.get("gdrive_secretaria")
        if p2:
            return Path(p2) / "Admissões"
        return Path.home() / "Documents"

    # ------------------------------------------------------------------
    # Dados
    # ------------------------------------------------------------------

    # ------------------------------------------------------------------
    # Dados combinados (residentes + candidatos)
    # ------------------------------------------------------------------

    def _todos(self) -> list[dict]:
        """Candidatos activos primeiro, depois residentes."""
        candidatos: list[dict] = []
        if hasattr(self.controller, "candidatos_repo"):
            try:
                for c in self.controller.candidatos_repo.get_all():
                    c["_tipo"]           = "candidato"
                    c["numero_residente"] = f"C{c['id']}"
                    c["especial"]        = c.get("estado", "pendente")
                    candidatos.append(c)
            except Exception:
                pass
        residentes = self.controller.residentes_repo.get_all()
        return candidatos + residentes

    def _carregar(self):
        self._popular(self._todos())
        # pré-selecção: candidato passado directamente
        if self._candidato is not None:
            cid = self._candidato.get("id")
            for iid, d in self._iid_map.items():
                if d.get("_tipo") == "candidato" and d.get("id") == cid:
                    self._tree.selection_set(iid)
                    self._tree.see(iid)
                    self._on_select()
                    break
        elif self._pre_numero is not None:
            for iid, d in self._iid_map.items():
                if d.get("numero_residente") == self._pre_numero:
                    self._tree.selection_set(iid)
                    self._tree.see(iid)
                    self._on_select()
                    break

    def _pesquisar(self):
        txt     = self._v_pesq.get().strip().lower()
        so_pend = self._v_so_pendentes.get()
        _ACTIVOS = {"pendente", "espera"}
        resultado = [
            d for d in self._todos()
            if (not txt or txt in (d.get("nome") or "").lower())
            and (not so_pend or str(d.get("especial") or "").lower() in _ACTIVOS)
        ]
        self._popular(resultado)

    def _popular(self, dados: list[dict]):
        self._tree.delete(*self._tree.get_children())
        self._iid_map.clear()
        so_pend  = self._v_so_pendentes.get()
        _ACTIVOS = {"pendente", "espera"}
        for d in dados:
            esp = str(d.get("especial") or "").strip().lower()
            if so_pend and esp not in _ACTIVOS:
                continue
            iid = self._tree.insert("", "end", values=(
                d.get("numero_residente", ""),
                d.get("nome", ""),
                d.get("especial", ""),
                d.get("data_admissao", ""),
            ))
            self._iid_map[iid] = d

    def _on_select(self, _event=None):
        sel = self._tree.selection()
        if not sel:
            self._dados = None
            self._lbl_ficha.config(text="Nenhum residente seleccionado.", fg="gray")
            self._lbl_preview.config(text="")
            return
        d = self._iid_map.get(sel[0])
        self._dados = d
        if d:
            linhas = [
                f"Nº {d.get('numero_residente')}  —  {d.get('nome', '')}",
                f"NIF: {d.get('nif') or '—'}   IBAN: {d.get('iban') or '—'}",
                f"Responsável: {d.get('responsavel') or '—'}   Tlm: {d.get('resp_tlm') or '—'}",
                f"Admissão: {d.get('data_admissao') or '—'}   Especial: {d.get('especial') or '—'}",
            ]
            self._lbl_ficha.config(text="\n".join(linhas), fg=self.FG)
            self._actualizar_preview()

    def _actualizar_preview(self):
        if self._dados:
            nf = self._nome_ficheiro(self._dados)
            pasta = self._v_destino.get().strip() or str(self._admissoes_dir())
            self._lbl_preview.config(
                text=f"Ficheiro a criar em  {pasta}:\n  • {nf}_contrato.docx",
                fg=self.FG,
            )
        else:
            self._lbl_preview.config(text="", fg="gray")

    def _nome_ficheiro(self, d: dict) -> str:
        nome = d.get("nome") or f"residente_{d.get('numero_residente', 'x')}"
        return "".join(c if c.isalnum() or c in " _-" else "_" for c in nome).strip()

    def _browse_destino(self):
        path = filedialog.askdirectory(
            title="Seleccionar pasta de destino",
            initialdir=self._v_destino.get() or str(self._admissoes_dir()),
            parent=self.root,
        )
        if path:
            self._v_destino.set(path)
            self._actualizar_preview()

    # ------------------------------------------------------------------
    # Geração
    # ------------------------------------------------------------------

    def _gerar(self):
        if not self._dados:
            messagebox.showwarning("Atenção", "Seleccione um residente.", parent=self.root)
            return

        pasta_dest = Path(self._v_destino.get().strip() or str(self._admissoes_dir()))
        nome_tmpl = self._v_minuta.get().strip()
        if not nome_tmpl:
            messagebox.showwarning("Atenção", "Seleccione uma minuta (template).", parent=self.root)
            return
        extras = getattr(self, "_minutas_extras", {})
        tmpl = extras.get(nome_tmpl) or (self._minutas_dir() / nome_tmpl)

        if not tmpl.exists():
            messagebox.showerror(
                "Minuta não encontrada",
                f"O ficheiro de minuta não foi encontrado:\n{tmpl}\n\n"
                f"Verifique a pasta:\n{self._minutas_dir()}",
                parent=self.root,
            )
            return

        try:
            from docx import Document
        except ImportError:
            messagebox.showerror(
                "Módulo em falta",
                "A biblioteca python-docx não está instalada.\nExecute: pip install python-docx",
                parent=self.root,
            )
            return

        pasta_dest.mkdir(parents=True, exist_ok=True)
        nf = self._nome_ficheiro(self._dados)
        destino = pasta_dest / f"{nf}_contrato.docx"

        mapping = _build_mapping(self._dados)
        try:
            doc = Document(tmpl)
            phs_doc      = _placeholders_em_doc(doc)
            em_falta     = sorted(ph for ph in phs_doc if ph in mapping and not mapping[ph])
            nao_mapeados = sorted(ph for ph in phs_doc if ph not in mapping)
            for ph in nao_mapeados:
                mapping[ph] = ""
            _replace_in_doc(doc, mapping)
            doc.save(destino)
            logger.info("Contrato gerado: %s", destino)
        except Exception as exc:
            logger.exception("Erro ao gerar contrato")
            messagebox.showerror("Erro", str(exc), parent=self.root)
            return

        if em_falta or nao_mapeados:
            messagebox.showwarning(
                "Campos sem valor",
                "Os seguintes campos ficam em branco no documento:\n\n" +
                "\n".join(f"  {ph}" for ph in em_falta + nao_mapeados),
                parent=self.root,
            )

        # oferecer actualizar estado após geração do contrato
        esp          = str(self._dados.get("especial") or "").lower()
        is_candidato = self._dados.get("_tipo") == "candidato"
        _ACTIVOS     = {"pendente", "espera"}
        if esp in _ACTIVOS:
            if is_candidato:
                pergunta = (f"O candidato «{self._dados.get('nome')}» está em estado «{esp}».\n"
                            "Deseja passar a «inscrito» agora que o contrato foi gerado?")
                titulo   = "Marcar como inscrito"
            else:
                pergunta = (f"O residente «{self._dados.get('nome')}» está marcado como «{esp}».\n"
                            "Deseja remover essa marcação agora que o contrato foi gerado?")
                titulo   = "Remover estado pendente"
            if messagebox.askyesno(titulo, pergunta, parent=self.root):
                try:
                    if is_candidato:
                        self.controller.candidatos_repo.update_estado(
                            int(self._dados["id"]), "inscrito"
                        )
                        self._dados["especial"] = "inscrito"
                    else:
                        self.controller.residentes_repo.update(
                            int(self._dados["numero_residente"]),
                            {"especial": None},
                        )
                        self._dados["especial"] = None
                except Exception as exc:
                    messagebox.showwarning("Aviso", f"Não foi possível actualizar: {exc}",
                                           parent=self.root)

        messagebox.showinfo(
            "Contrato gerado",
            f"Ficheiro guardado em:\n{pasta_dest}\n\n• {destino.name}",
            parent=self.root,
        )
